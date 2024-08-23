import mimetypes
from textwrap import wrap
from django.core.files.storage import FileSystemStorage
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.models import User
from django.contrib import messages
from .models import UserProfile
from .models import UploadedFile
from .forms import FileUploadForm
from pdf2image import convert_from_path
import os
import io

os.environ['TESSDATA_PREFIX'] = 'D:/AppGallery/Tesseract-OCR/tessdata/'
import pytesseract

pytesseract.pytesseract.tesseract_cmd = r"D:\AppGallery\Tesseract-OCR\tesseract.exe"
from pdf2image import convert_from_path
import docx
from django.http import JsonResponse
from django.urls import reverse
from django.conf import settings
from PIL import Image
import logging

logger = logging.getLogger(__name__)  # 使用全局配置过的logger
import time
from googletrans import Translator
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import json
from reportlab.lib.pagesizes import A4
from reportlab.lib import fonts
import requests
import hashlib
import random
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from django.core.files.storage import default_storage
from transformers import pipeline
from datetime import datetime
from transformers import MarianMTModel, MarianTokenizer
import sentencepiece
import sacremoses
from django.contrib.auth.models import AnonymousUser
from django.utils.deprecation import MiddlewareMixin


# Create your views here.
def get_files_info(directory):
    files_info = []
    directory_path = os.path.join(settings.BASE_DIR, directory)
    for file_name in os.listdir(directory_path):
        file_path = os.path.join(directory_path, file_name)
        if os.path.isfile(file_path):
            file_info = {
                'name': file_name,
                'path': os.path.relpath(file_path, settings.BASE_DIR),
                'last_modified': datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d'),
                'size': f"{os.path.getsize(file_path) / (1024 * 1024):.2f} MB",
            }
            files_info.append(file_info)
    return files_info


def translate_text_function(text):
    # 本地模型路径
    model_path = "media/opus-mt-zh-en"

    # 从本地路径加载模型和分词器
    model = MarianMTModel.from_pretrained(model_path)
    tokenizer = MarianTokenizer.from_pretrained(model_path)

    # 将输入文本进行分词
    inputs = tokenizer(text, return_tensors="pt", padding=True)

    # 使用模型进行翻译
    translated = model.generate(
        **inputs,
        num_beams=10,  # 使用5个束搜索
        early_stopping=True  # 所有束生成结束符时停止生成
    )

    # 解码翻译后的文本
    translated_text = tokenizer.decode(translated[0], skip_special_tokens=True)

    return translated_text


# 登录视图
def login_view(request):
    if request.method == "POST":
        phone_number = request.POST.get('username')
        password = request.POST.get('password')

        # 验证用户信息
        try:
            user_profile = UserProfile.objects.get(phone_number=phone_number)
            if user_profile.password == password:  # 假设密码以明文形式存储
                # 登录成功，保存用户 session 或执行其他登录逻辑
                request.session['user_id'] = user_profile.id  # 在 session 中存储用户ID
                messages.success(request, "登录成功！")
                return redirect('/my_files/')
            else:
                messages.error(request, "密码错误，请重试。")
        except UserProfile.DoesNotExist:
            messages.error(request, "该手机号未注册，请先注册。")

    return render(request, 'login.html')


# 注册视图
def register_view(request):
    if request.method == "POST":
        phone_email = request.POST.get('phone-email')
        password = request.POST.get('password')
        agree = request.POST.get('agree')

        if not agree:
            messages.error(request, "您必须同意服务协议和隐私政策才能注册。")
        else:
            # 检查是否手机号/邮箱已存在
            if UserProfile.objects.filter(phone_number=phone_email).exists():
                messages.error(request, "手机号/邮箱已被注册，请使用其他手机号/邮箱。")
            else:
                # 创建新的用户
                UserProfile.objects.create(phone_number=phone_email, password=password)
                messages.success(request, "注册成功，请登录。")
                return redirect('login')

    return render(request, 'register.html')


def login_first(request):
    return render(request, 'login_first.html')


def help_center(request):
    return render(request, 'help_center.html')


# 文件管理视图
def my_files(request):
    user_profile_id = request.session.get('user_id')
    if not user_profile_id:
        messages.error(request, "请先登录。")
        return redirect('/login/')

    user_profile = UserProfile.objects.get(id=user_profile_id)

    if request.method == "POST":
        form = FileUploadForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = form.save(commit=False)
            uploaded_file.user_profile = user_profile  # 关联当前用户
            mime_type, _ = mimetypes.guess_type(uploaded_file.file.name)
            uploaded_file.attributes = mime_type or 'Unknown'
            uploaded_file.save()
            return redirect('my_files')  # 上传成功后重新加载页面
    else:
        form = FileUploadForm()

    # 搜索功能
    search_query = request.GET.get('search', '')

    # 排序功能
    sort_by = request.GET.get('sort_by', 'default')

    # 获取当前用户的所有文件，并默认按上传时间排序
    files = UploadedFile.objects.filter(user_profile=user_profile).order_by('-upload_date')

    if search_query:
        files = files.filter(name__icontains=search_query)

    if sort_by == 'name':
        files = files.order_by('name')
    elif sort_by == 'recent':
        files = files.order_by('-upload_date')
    else:
        files = files.order_by('-id')  # 默认排序为按ID排序

    context = {
        'form': form,
        'files': files,
        'search_query': search_query,
        'sort_by': sort_by,
    }

    return render(request, 'my_files.html', context)


logger = logging.getLogger(__name__)


def text_recognition(request):
    form = FileUploadForm()
    uploaded_file = UploadedFile.objects.last()  # 只显示最后一次上传的文件
    ocr_text = None
    pdf_url = None  # 在这里定义 pdf_url 的初始值
    document_text = "这里是显示在页面中央的文档内容"

    if uploaded_file:
        pdf_url = uploaded_file.file.url  # 确保 pdf_url 在这里被赋值
        logger.info(f"Processing file: {uploaded_file.file.path}")
        file_extension = os.path.splitext(uploaded_file.file.path)[1].lower()

        if file_extension == '.pdf':
            try:
                ocr_text = extract_text_from_pdf(uploaded_file.file.path)
                logger.info(f"OCR 处理成功，文件路径: {uploaded_file.file.path}")
            except Exception as e:
                logger.error(f"OCR 处理失败: {e}")
        else:
            logger.error(f"文件类型不支持: {file_extension}")

    # 确保在这里 pdf_url 已经被正确赋值
    if pdf_url:
        document_text = pdf_url  # 将 PDF 的 URL 赋值给 document_text

    context = {
        'form': form,
        'uploaded_file': uploaded_file,
        'ocr_text': ocr_text,
        'pdf_url': pdf_url,
        'document_text': document_text  # 将 PDF 的 URL 传递到模板中
    }
    return render(request, 'text_recognition.html', context)


def upload_file(request):
    if request.method == 'POST':
        form = FileUploadForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES.get('file')
            if uploaded_file.content_type == 'application/pdf':
                form.save()
                logger.info(f"文件上传成功: {uploaded_file.name}")
            else:
                logger.error(f"非法文件类型上传: {uploaded_file.content_type}")
                return HttpResponse("仅支持 PDF 文件", status=400)
            return redirect(reverse('text_recognition'))
    return redirect(reverse('text_recognition'))


def delete_file(request, file_id):
    uploaded_file = UploadedFile.objects.get(id=file_id)
    if os.path.exists(uploaded_file.file.path):
        os.remove(uploaded_file.file.path)
    uploaded_file.delete()
    return redirect(reverse('text_recognition'))


def generate_doc(request):
    if request.method == 'POST':
        text = request.POST.get('text')
        doc = docx.Document()
        doc.add_paragraph(text)
        doc_path = os.path.join(settings.MEDIA_ROOT, 'download', 'recognized_text.docx')
        doc.save(doc_path)
        response = HttpResponse(open(doc_path, 'rb').read(),
                                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response['Content-Disposition'] = f'attachment; filename=recognized_text.docx'
        os.remove(doc_path)
        return response


def extract_text_from_pdf(pdf_path):
    start_time = time.time()  # 开始时间

    images = convert_from_path(pdf_path)
    text = ''
    for img in images:
        text += pytesseract.image_to_string(img, lang='chi_sim')  # 使用中文语言包

    end_time = time.time()  # 结束时间
    logger.info(f"OCR处理时间: {end_time - start_time} 秒")  # 输出处理时间
    return text


def translate_text_view(request):
    if request.method == 'POST':
        original_text = request.POST.get('text')
        max_length = 1000  # 假设字数限制为1000字

        if len(original_text) > max_length:
            # 如果字数超限，返回错误信息
            return JsonResponse({'error': '太多字了啊'}, status=400)

        translated_text = translate_text_function(original_text)  # 替换为你的翻译逻辑

        return JsonResponse({'translated_text': translated_text})


def export_as_pdf(request):
    if request.method == 'POST':
        text = request.POST.get('text')
        doc = docx.Document()
        doc.add_paragraph(text)

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="document.pdf"'

        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)

        # 注册SimSun中文字体，确保路径正确
        pdfmetrics.registerFont(TTFont('SimSun', 'C://Windows//Fonts//simsunb.ttf'))

        # 获取文档中的文本内容
        document_text = "\n".join([para.text for para in doc.paragraphs])

        # 设置字体和大小
        p.setFont("SimSun", 12)

        # 设置文本对象起始位置
        text_object = p.beginText(100, 750)

        # 最大行宽和初始高度
        max_width = 800  # 每行的最大字符宽度
        height = 750

        for line in document_text.split("\n"):
            # 使用textwrap处理自动换行，这里假设每个中文字符占一个字符宽度
            wrapped_lines = wrap(line, max_width // 12)
            for wrapped_line in wrapped_lines:
                text_object.textLine(wrapped_line)
                height -= 14  # 每行文本降低的高度
                if height < 40:  # 当接近页面底部时，开始新页
                    p.drawText(text_object)
                    p.showPage()
                    text_object = p.beginText(100, 750)
                    text_object.setFont("SimSun", 12)
                    height = 750

        # 绘制文本到页面
        p.drawText(text_object)

        # 完成PDF的生成
        p.showPage()
        p.save()

        # 将PDF内容写入响应
        buffer.seek(0)
        response.write(buffer.getvalue())
        buffer.close()

        return response


logger = logging.getLogger(__name__)


def text_classification(request):
    user_profile_id = request.session.get('user_id')
    if not user_profile_id:
        messages.error(request, "请先登录。")
        return redirect('/login/')

    user_profile = UserProfile.objects.get(id=user_profile_id)

    form = FileUploadForm()
    uploaded_file = UploadedFile.objects.filter(user_profile=user_profile).last()  # 只显示该用户最后一次上传的文件
    success_message = None  # 初始化提示信息

    # 获取用户专属目录中的所有文件
    uploads_path = os.path.join(settings.MEDIA_ROOT, 'uploads', user_profile.phone_number)
    files_info = []

    if os.path.exists(uploads_path):
        for file_name in os.listdir(uploads_path):
            file_path = os.path.join(uploads_path, file_name)

            if os.path.isfile(file_path):
                file_info = {
                    'name': file_name,
                    'path': os.path.relpath(file_path, settings.BASE_DIR),
                    'last_modified': datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d'),
                    'size': f"{os.path.getsize(file_path) / (1024 * 1024):.2f} MB",
                }
                files_info.append(file_info)

    # 分类标签与路径映射，路径按用户区分
    label_to_path = {
        '学习文档_教材': f'{uploads_path}/learning_document/jiao_cai/',
        '学习文档_试卷': f'{uploads_path}/learning_document/shi_juan/',
        '学习文档_笔记': f'{uploads_path}/learning_document/bi_ji/',
        '生活文档_食谱': f'{uploads_path}/life_document/shi_pu/',
        '生活文档_攻略': f'{uploads_path}/life_document/gong_lue/',
        '生活文档_计划': f'{uploads_path}/life_document/ji_hua/',
        '工作文档_财务': f'{uploads_path}/work_document/cai_wu/',
        '工作文档_方案': f'{uploads_path}/work_document/fang_an/',
        '工作文档_合同': f'{uploads_path}/work_document/he_tong/',
    }

    if uploaded_file:
        try:
            ocr_text = extract_text_from_pdf(uploaded_file.file.path)
            classifier = pipeline("zero-shot-classification", model="media/chinese-roberta-wwm-ext")
            candidate_labels = list(label_to_path.keys())
            result = classifier(ocr_text, candidate_labels)
            classification = result['labels'][0]

            pdf_filename = os.path.basename(uploaded_file.file.path)
            save_path = label_to_path.get(classification)

            if save_path:
                os.makedirs(save_path, exist_ok=True)  # 如果目录不存在则创建
                full_file_path = os.path.join(save_path, pdf_filename)
                with open(uploaded_file.file.path, 'rb') as f:
                    pdf_content = f.read()
                with open(full_file_path, 'wb') as file:
                    file.write(pdf_content)
                success_message = f"文件已成功存入 {classification} 文件夹。"
            else:
                success_message = "没有找到匹配的标签路径，无法保存文件。"

        except Exception as e:
            logger.error(f"OCR 处理失败: {e}")
            success_message = "文件分类失败，请重试。"

    return render(request, 'text_classification.html',
                  {'files_info': files_info, 'form': form, 'success_message': success_message})


def create_directory_if_not_exists(directory):
    if not os.path.exists(directory):
        os.makedirs(directory)

def get_files_info(directory):
    files_info = []
    if os.path.exists(directory):
        for file_name in os.listdir(directory):
            file_path = os.path.join(directory, file_name)
            if os.path.isfile(file_path):
                file_info = {
                    'name': file_name,
                    'path': os.path.relpath(file_path, settings.BASE_DIR),
                    'last_modified': datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d'),
                    'size': f"{os.path.getsize(file_path) / (1024 * 1024):.2f} MB",
                }
                files_info.append(file_info)
    return files_info

def work_document(request):
    user_profile_id = request.session.get('user_id')
    if not user_profile_id:
        messages.error(request, "请先登录。")
        return redirect('/login/')

    user_profile = UserProfile.objects.get(id=user_profile_id)

    directories = [
        f'media/uploads/{user_profile.phone_number}/work_document/cai_wu/',
        f'media/uploads/{user_profile.phone_number}/work_document/fang_an/',
        f'media/uploads/{user_profile.phone_number}/work_document/he_tong/'
    ]
    all_files_info = []
    for directory in directories:
        create_directory_if_not_exists(directory)
        all_files_info.extend(get_files_info(directory))

    return render(request, 'work_document.html', {'files_info': all_files_info})

def finance_category(request):
    user_profile_id = request.session.get('user_id')
    if not user_profile_id:
        messages.error(request, "请先登录。")
        return redirect('/login/')

    user_profile = UserProfile.objects.get(id=user_profile_id)
    directory = f'media/uploads/{user_profile.phone_number}/work_document/cai_wu/'
    create_directory_if_not_exists(directory)
    files_info = get_files_info(directory)
    return render(request, 'finance_category.html', {'files_info': files_info})

def plan_category(request):
    user_profile_id = request.session.get('user_id')
    if not user_profile_id:
        messages.error(request, "请先登录。")
        return redirect('/login/')

    user_profile = UserProfile.objects.get(id=user_profile_id)
    directory = f'media/uploads/{user_profile.phone_number}/work_document/fang_an/'
    create_directory_if_not_exists(directory)
    files_info = get_files_info(directory)
    return render(request, 'plan_category.html', {'files_info': files_info})

def contract_category(request):
    user_profile_id = request.session.get('user_id')
    if not user_profile_id:
        messages.error(request, "请先登录。")
        return redirect('/login/')

    user_profile = UserProfile.objects.get(id=user_profile_id)
    directory = f'media/uploads/{user_profile.phone_number}/work_document/he_tong/'
    create_directory_if_not_exists(directory)
    files_info = get_files_info(directory)
    return render(request, 'contract_category.html', {'files_info': files_info})

def learning_document(request):
    user_profile_id = request.session.get('user_id')
    if not user_profile_id:
        messages.error(request, "请先登录。")
        return redirect('/login/')

    user_profile = UserProfile.objects.get(id=user_profile_id)

    directories = [
        f'media/uploads/{user_profile.phone_number}/learning_document/bi_ji/',
        f'media/uploads/{user_profile.phone_number}/learning_document/jiao_cai/',
        f'media/uploads/{user_profile.phone_number}/learning_document/shi_juan/'
    ]
    all_files_info = []
    for directory in directories:
        create_directory_if_not_exists(directory)
        all_files_info.extend(get_files_info(directory))

    return render(request, 'learning_document.html', {'files_info': all_files_info})

def note_category(request):
    user_profile_id = request.session.get('user_id')
    if not user_profile_id:
        messages.error(request, "请先登录。")
        return redirect('/login/')

    user_profile = UserProfile.objects.get(id=user_profile_id)
    directory = f'media/uploads/{user_profile.phone_number}/learning_document/bi_ji/'
    create_directory_if_not_exists(directory)
    files_info = get_files_info(directory)
    return render(request, 'note_category.html', {'files_info': files_info})

def textbook_category(request):
    user_profile_id = request.session.get('user_id')
    if not user_profile_id:
        messages.error(request, "请先登录。")
        return redirect('/login/')

    user_profile = UserProfile.objects.get(id=user_profile_id)
    directory = f'media/uploads/{user_profile.phone_number}/learning_document/jiao_cai/'
    create_directory_if_not_exists(directory)
    files_info = get_files_info(directory)
    return render(request, 'textbook_category.html', {'files_info': files_info})

def testpaper_category(request):
    user_profile_id = request.session.get('user_id')
    if not user_profile_id:
        messages.error(request, "请先登录。")
        return redirect('/login/')

    user_profile = UserProfile.objects.get(id=user_profile_id)
    directory = f'media/uploads/{user_profile.phone_number}/learning_document/shi_juan/'
    create_directory_if_not_exists(directory)
    files_info = get_files_info(directory)
    return render(request, 'testpaper_category.html', {'files_info': files_info})

def life_document(request):
    user_profile_id = request.session.get('user_id')
    if not user_profile_id:
        messages.error(request, "请先登录。")
        return redirect('/login/')

    user_profile = UserProfile.objects.get(id=user_profile_id)

    directories = [
        f'media/uploads/{user_profile.phone_number}/life_document/gong_lue/',
        f'media/uploads/{user_profile.phone_number}/life_document/ji_hua/',
        f'media/uploads/{user_profile.phone_number}/life_document/shi_pu/'
    ]
    all_files_info = []
    for directory in directories:
        create_directory_if_not_exists(directory)
        all_files_info.extend(get_files_info(directory))

    return render(request, 'life_document.html', {'files_info': all_files_info})

def introduction_category(request):
    user_profile_id = request.session.get('user_id')
    if not user_profile_id:
        messages.error(request, "请先登录。")
        return redirect('/login/')

    user_profile = UserProfile.objects.get(id=user_profile_id)
    directory = f'media/uploads/{user_profile.phone_number}/life_document/gong_lue/'
    create_directory_if_not_exists(directory)
    files_info = get_files_info(directory)
    return render(request, 'introduction_category.html', {'files_info': files_info})

def scheme_category(request):
    user_profile_id = request.session.get('user_id')
    if not user_profile_id:
        messages.error(request, "请先登录。")
        return redirect('/login/')

    user_profile = UserProfile.objects.get(id=user_profile_id)
    directory = f'media/uploads/{user_profile.phone_number}/life_document/ji_hua/'
    create_directory_if_not_exists(directory)
    files_info = get_files_info(directory)
    return render(request, 'scheme_category.html', {'files_info': files_info})

def cookbook_category(request):
    user_profile_id = request.session.get('user_id')
    if not user_profile_id:
        messages.error(request, "请先登录。")
        return redirect('/login/')

    user_profile = UserProfile.objects.get(id=user_profile_id)
    directory = f'media/uploads/{user_profile.phone_number}/life_document/shi_pu/'
    create_directory_if_not_exists(directory)
    files_info = get_files_info(directory)
    return render(request, 'cookbook_category.html', {'files_info': files_info})


def more_search(request):
    return render(request, 'more_search.html')


def set_up(request):
    return render(request, 'set_up.html')
